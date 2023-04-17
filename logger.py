"""
Скрипт для логирования и записи данных лога в файл.
Также определения времени выполнения команды.
"""

import csv
import datetime
import os
import http.client  # Для определения своего WAN адреса
import socket  # Для определения своего LAN адреса

import docx

# Release v1.7.1
RELEASE = "v1.7.1"
OBJECT_NAME = "UNKNOWN"


def object_name():
    global OBJECT_NAME

    file_path = "config/name_object.txt"

    if os.path.exists(file_path):
        with open(file_path, "r", encoding='UTF-8') as file:
            OBJECT_NAME = file.readline()

            if len(OBJECT_NAME) == 0:
                OBJECT_NAME = input("Введите имя объекта:\n")
                file.write(OBJECT_NAME)

    else:
        with open(file_path, "x", encoding='UTF-8') as file:
            OBJECT_NAME = input("Введите имя объекта:\n")
            file.write(OBJECT_NAME)

    return OBJECT_NAME


def get_time(format: str = "time") -> str:
    local_date = datetime.datetime.now().date()
    local_time = datetime.datetime.now().time()
    gmt_date = datetime.datetime.utcnow().date()
    gmt_time = datetime.datetime.utcnow().time()

    if format == "time":
        time_format = '%H:%M:%S'
        return f"{local_time.strftime(time_format)} (GMT {gmt_time.strftime(time_format)})"

    elif format == "date":
        date_format = '%Y-%m-%d'
        return f"DATE {local_date.strftime(date_format)} (GMT {date_format})"

    elif format == "for_csv":
        return f"{gmt_date}_{gmt_time.strftime('%H%M%S')}_GMT"

    elif format == "for_docx":
        return local_date.strftime('%d%m%y')

    elif format == "for_pu":
        date_format = '%y%m%d'
        gmt_tomorrow = gmt_date + datetime.timedelta(days=1)

        return f"{gmt_date.strftime(date_format)}000000Z\n{gmt_tomorrow.strftime(date_format)}000000Z"


def my_wan_ip():
    # Определение моего WAN адреса
    wan_ip = http.client.HTTPConnection("ifconfig.me")
    wan_ip.request("GET", "/ip")
    return wan_ip.getresponse().read().decode('utf-8')


def my_lan_ip():
    # Определение моего LAN адреса
    return socket.gethostbyname(socket.gethostname())


def file_for_csv():
    # Функция создания лог файла ГГММДД_ччммсс_GMT.log

    global OBJECT_NAME

    # Создаём директорию logs
    try:
        os.mkdir("logs")
    except FileExistsError:
        pass

    # Очищаем папку logs оставляя максимум 10 файлов
    os.chdir("logs")  # Меняем рабочую директорию
    files_in_dir_logs = tuple(os.walk(os.getcwd()))[0][-1]  # Получаем список файлов внутри ./logs
    for i in range(len(files_in_dir_logs) - 10):
        os.remove(files_in_dir_logs[i])
    os.chdir("../")  # Меняем рабочую директорию

    # Создаём директорию logs_in_docx
    try:
        os.mkdir("logs_in_docx")
    except FileExistsError:
        pass

    # Очищаем папку logs_in_docx оставляя максимум 10 файлов
    os.chdir("logs_in_docx")  # Меняем рабочую директорию
    files_in_dir_logs = tuple(os.walk(os.getcwd()))[0][-1]  # Получаем список файлов внутри ./logs_in_docx
    for i in range(len(files_in_dir_logs) - 10):
        os.remove(files_in_dir_logs[i])
    os.chdir("../")  # Меняем рабочую директорию

    os.chdir("logs")  # Меняем рабочую директорию
    logs_file_name = f"{OBJECT_NAME} {get_time('for_csv')}.csv"
    logs_file = open(logs_file_name, mode="x", encoding="utf-8")  # Создаём и открываем файл в режиме записи
    logs_file.write("protocol;time;resource;size;from;to;msg;error;\n")
    logs_file.write(f"{OBJECT_NAME};{get_time()};WAN {my_wan_ip()};LAN {my_lan_ip()};;;;;")
    logs_file.close()  # Закрываем файл
    os.chdir("../")  # Меняем рабочую директорию


def log_csv(text):
    # Функция для записи данных в файл ГГММДД_ччммсс_GMT.log

    os.chdir("logs")  # Меняем рабочую директорию
    name_file = tuple(os.walk(os.getcwd()))[0][-1]  # Получаем список файлов внутри ./logs
    name_file.sort()
    logs_file = open(name_file[-1], mode="a", encoding="utf-8")  # Открываем файл в режиме дозаписи
    logs_file.write("\n" + text)  # Дозаписываем в файл
    logs_file.close()  # Закрываем файл
    os.chdir("../")  # Меняем рабочую директорию


def csv_to_docx() -> str:
    os.chdir("logs")
    name_file = tuple(os.walk(os.getcwd()))[0][-1]
    name_file.sort()
    with open(name_file[-1], "r", encoding="utf-8") as log_in_csv:
        log_list = [line[0].split(";") for line in csv.reader(log_in_csv)]
        log_in_csv.close()
    os.chdir("../")

    http_time_list = []
    email_time_list = []
    im_time_list = []
    voip_time_list = []
    ftp_time_list = []
    telnet_time_list = []
    ssh_time_list = []
    https_time_list = []

    for test in log_list:
        if test[0] == "HTTP":
            http_time_list.append(test[1][:5])
        elif "EMAIL" in test[0]:
            email_time_list.append(test[1][:5])
        elif "IM" in test[0]:
            im_time_list.append(test[1][:5])
        elif "VOIP" in test[0]:
            voip_time_list.append(test[1][:5])
        elif "FTP" in test[0]:
            if test[3] == "0":  # start
                ftp_time_list.append(test[1][:5])
            else:  # end
                # Получаем строку "start - hh:mi (size MB/GB)"
                ftp_time_list[-1] += f" - {test[1][:5]} ({test[3].replace(' (', '    ')[:8].strip(' ')})"
        elif "TELNET" in test[0]:
            telnet_time_list.append(test[1][:5])
        elif "SSH" in test[0]:
            ssh_time_list.append(test[1][:5])
        elif "HTTPS" in test[0]:
            https_time_list.append(test[1][:5])

    # VOIP тест не готов, поэтому заменим его первым временем HTTP,
    # т.к. он делается вручную перед запуском авто-теста
    try:
        if len(voip_time_list) == 0:
            voip_time_list.append(http_time_list[0])
            voip_time_list.append(http_time_list[0])
    except IndexError:
        pass

    wordDoc = docx.Document("data/template_report_test.docx")  # Открываем чистую таблицу
    style = wordDoc.styles['Normal']  # задаем стиль текста по умолчанию
    style.font.name = 'Times New Roman'  # название шрифта
    style.font.size = docx.shared.Pt(12)  # размер шрифта

    # Таблица №1 wordDoc.tables[0]
    wordDoc.tables[0].rows[0].cells[0].text = f"Оператор: {OBJECT_NAME}"
    wordDoc.tables[0].rows[2].cells[0].text = f"IP-адрес тестового рабочего места: {log_list[1][2][4:]}"
    wordDoc.tables[0].rows[3].cells[0].text = f"Дата выполнения тестов: {get_time('date')}"

    # Таблица №2 wordDoc.tables[1]
    try:
        # HTTP
        wordDoc.tables[1].rows[1].cells[3].text = http_time_list[0]  # HTTP Время http://kremlin.ru
        wordDoc.tables[1].rows[2].cells[3].text = http_time_list[1]  # HTTP Время http://kremlin.ru/structure/president
        wordDoc.tables[1].rows[3].cells[3].text = http_time_list[2]  # HTTP Время http://kremlin.ru/static/pdf/*.pdf
        wordDoc.tables[1].rows[4].cells[3].text = http_time_list[3]  # HTTP Время http://fsb.ru
        wordDoc.tables[1].rows[5].cells[3].text = http_time_list[4]  # HTTP Время http://rtc-nt.ru
        wordDoc.tables[1].rows[6].cells[3].text = http_time_list[5]  # HTTP Время http://gramota.ru
        wordDoc.tables[1].rows[7].cells[3].text = http_time_list[6]  # HTTP Время http://duma.gov.ru
        wordDoc.tables[1].rows[8].cells[3].text = http_time_list[7]  # HTTP Время http://ivo.garant.ru
        wordDoc.tables[1].rows[9].cells[3].text = http_time_list[8]  # HTTP Время http://www.thesheep.info
        wordDoc.tables[1].rows[10].cells[3].text = http_time_list[9]  # HTTP Время http://www.grani.ru
    except IndexError:
        pass

    try:
        # EMAIL
        wordDoc.tables[1].rows[11].cells[3].text = email_time_list[
            0]  # EMAIL Время Отправка письма с 3 получателями, вложением и копией
        wordDoc.tables[1].rows[13].cells[3].text = email_time_list[
            1]  # EMAIL Время Получение письма с 3 получателями и вложением
        wordDoc.tables[1].rows[15].cells[3].text = email_time_list[
            2]  # EMAIL Время Отправка письма с 2 копиями и иероглифы
        wordDoc.tables[1].rows[17].cells[3].text = email_time_list[
            3]  # EMAIL Время Получение письма с 2 копиями и иероглифы
    except IndexError:
        pass

    try:
        # IM
        wordDoc.tables[1].rows[19].cells[3].text = im_time_list[2]  # IM Время Отправка сообщений
        wordDoc.tables[1].rows[20].cells[3].text = im_time_list[3]  # IM Время Получение сообщений
    except IndexError:
        pass

    try:
        # VOIP
        wordDoc.tables[1].rows[21].cells[3].text = voip_time_list[0]  # VOIP Время Исходящее голосовое соединение
        wordDoc.tables[1].rows[23].cells[3].text = voip_time_list[1]  # VOIP Время Входящее голосовое соединение
    except IndexError:
        pass

    try:
        # FTP
        wordDoc.tables[1].rows[25].cells[3].text = ftp_time_list[0][:5]  # FTP Время Доступ к ресурсу
        wordDoc.tables[1].rows[26].cells[3].text = "\n" + ftp_time_list[0]  # FTP Время start - end ts.zip
        wordDoc.tables[1].rows[27].cells[3].text = ftp_time_list[1]  # FTP Время start - end  gtdw.zip
        wordDoc.tables[1].rows[28].cells[3].text = ftp_time_list[2]  # FTP Время start - end  maximum.zip
    except IndexError:
        pass

    try:
        # TELNET
        wordDoc.tables[1].rows[29].cells[3].text = telnet_time_list[0]  # TELNET Время 54.39.129.129
        wordDoc.tables[1].rows[30].cells[3].text = telnet_time_list[1]  # TELNET Время 192.241.222.161
        wordDoc.tables[1].rows[31].cells[3].text = telnet_time_list[2]  # TELNET Время 35.185.12.150
    except IndexError:
        pass

    try:
        # SSH
        wordDoc.tables[1].rows[32].cells[3].text = ssh_time_list[0]  # SSH Время 195.144.107.198
        wordDoc.tables[1].rows[33].cells[3].text = ssh_time_list[1]  # SSH Время 205.166.94.16
    except IndexError:
        pass

    try:
        # HTTPS
        wordDoc.tables[1].rows[34].cells[3].text = https_time_list[0]  # HTTPS Время https://yandex.ru
        wordDoc.tables[1].rows[35].cells[3].text = https_time_list[1]  # HTTPS Время https://mail.ru
        wordDoc.tables[1].rows[36].cells[3].text = https_time_list[2]  # HTTPS Время https://rambler.ru
        wordDoc.tables[1].rows[37].cells[3].text = https_time_list[3]  # HTTPS Время https://2ip.ru
        wordDoc.tables[1].rows[38].cells[3].text = https_time_list[4]  # HTTPS Время https://cia.gov
        wordDoc.tables[1].rows[39].cells[3].text = https_time_list[5]  # HTTPS Время https://nsa.gov
        wordDoc.tables[1].rows[40].cells[3].text = https_time_list[6]  # HTTPS Время https://ssu.gov.ua
        wordDoc.tables[1].rows[41].cells[3].text = https_time_list[7]  # HTTPS Время https://mossad.gov.il
        wordDoc.tables[1].rows[42].cells[3].text = https_time_list[8]  # HTTPS Время https://sis.gov.uk
        wordDoc.tables[1].rows[43].cells[3].text = https_time_list[9]  # HTTPS Время https://bnd.bund.de
    except IndexError:
        pass

    docx_file_name = f"{OBJECT_NAME} {get_time('for_docx')} Тесты ПСИ 573 ПД.docx"
    wordDoc.save(f"./logs_in_docx/{docx_file_name}")

    return docx_file_name
